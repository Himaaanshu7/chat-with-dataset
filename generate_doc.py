"""Generate Chat_With_Dataset_Documentation.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x45, 0x8E)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def divider():
    doc.add_paragraph("─" * 80)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data, row in zip(rows, table.rows[1:]):
        for i, val in enumerate(row_data):
            row.cells[i].text = val
    return table

# ── Cover Page ────────────────────────────────────────────────────────────────
title = doc.add_heading("Chat With Your Dataset", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x45, 0x8E)

sub = doc.add_paragraph("Project Documentation — AI-Powered Data Analysis Application")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(13)

meta = doc.add_paragraph(
    f"Author: Himanshu Mishra\n"
    f"Date: {datetime.date.today().strftime('%B %d, %Y')}\n"
    f"Live URL: https://chat-with-dataset-w3q3tiwx9b8duna8bdw9u6.streamlit.app\n"
    f"GitHub: https://github.com/Himaaanshu7/chat-with-dataset"
)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(11)

doc.add_page_break()

# ── 1. Executive Summary ──────────────────────────────────────────────────────
h1("1. Executive Summary")
body(
    "Chat With Your Dataset is a production-grade AI web application that enables users to "
    "upload any CSV or Excel file and interrogate it using plain English questions. "
    "The system leverages a large language model (Llama 3.3 70B via Groq) to translate "
    "natural language into executable Pandas code, runs the analysis safely, auto-generates "
    "appropriate visualizations, and returns plain-English explanations with suggested "
    "follow-up questions — creating an experience comparable to having a personal data analyst."
)
doc.add_paragraph()
body("The application is deployed publicly on Streamlit Community Cloud and is suitable for:")
bullet("Business analysts who need to explore datasets without writing code")
bullet("Data scientists who want rapid EDA (Exploratory Data Analysis)")
bullet("Students learning data analysis concepts")
bullet("Product managers reviewing business metrics")
bullet("Recruiters evaluating AI/data engineering portfolios")

doc.add_page_break()

# ── 2. Use Cases ──────────────────────────────────────────────────────────────
h1("2. Use Cases")

h2("2.1 Business Intelligence & Reporting")
add_table(
    ["Use Case", "Example Question", "Output"],
    [
        ["Revenue Analysis", "What are the top 5 products by revenue?", "Ranked table + bar chart"],
        ["Trend Detection", "Show monthly sales trend over the past year", "Line chart + explanation"],
        ["Regional Comparison", "Which region has the highest average order value?", "Table + bar chart"],
        ["KPI Monitoring", "What is the total revenue for Q3?", "Single metric value"],
        ["Drop Analysis", "Why did sales drop in March?", "Analysis + AI explanation"],
    ]
)

doc.add_paragraph()
h2("2.2 Exploratory Data Analysis (EDA)")
add_table(
    ["Use Case", "Example Question", "Output"],
    [
        ["Data Profiling", "Give me a summary of this dataset", "Stats table"],
        ["Distribution Analysis", "Show distribution of customer ages", "Histogram"],
        ["Outlier Detection", "Are there any outliers in the price column?", "Statistical summary"],
        ["Correlation Study", "Is there a correlation between price and quantity?", "Scatter plot + r-value"],
        ["Missing Data", "Which columns have the most missing values?", "Table"],
    ]
)

doc.add_paragraph()
h2("2.3 Customer & User Analytics")
add_table(
    ["Use Case", "Example Question", "Output"],
    [
        ["Top Customers", "Who are the top 10 customers by lifetime value?", "Ranked table"],
        ["Segmentation", "How many customers are in each segment?", "Bar chart"],
        ["Churn Analysis", "What percentage of customers churned last month?", "Metric"],
        ["Cohort View", "Show new vs returning customers by month", "Line chart"],
    ]
)

doc.add_paragraph()
h2("2.4 HR & People Analytics")
add_table(
    ["Use Case", "Example Question", "Output"],
    [
        ["Salary Analysis", "What is the average salary by department?", "Bar chart"],
        ["Headcount", "How many employees joined each year?", "Line chart"],
        ["Attrition", "Which department has the highest turnover rate?", "Table"],
        ["Demographics", "Show age distribution of employees", "Histogram"],
    ]
)

doc.add_paragraph()
h2("2.5 E-Commerce & Retail")
add_table(
    ["Use Case", "Example Question", "Output"],
    [
        ["Product Performance", "Which product category has the most returns?", "Bar chart"],
        ["Inventory Analysis", "Which products are low on stock?", "Table"],
        ["Pricing Study", "Show price vs sales volume scatter", "Scatter plot"],
        ["Seasonal Trends", "Show revenue by month for each category", "Line chart"],
    ]
)

doc.add_page_break()

# ── 3. Functional Requirements ────────────────────────────────────────────────
h1("3. Functional Requirements")

h2("3.1 File Upload")
bullet("FR-01: ", "System shall accept CSV and Excel (.xlsx, .xls) file formats")
bullet("FR-02: ", "System shall auto-detect and parse datetime columns")
bullet("FR-03: ", "System shall display a dataset summary upon upload (rows, columns, nulls, duplicates)")
bullet("FR-04: ", "System shall show column names, data types, and null percentages")
bullet("FR-05: ", "System shall display a sample of the first 5 rows")
bullet("FR-06: ", "System shall compute and display descriptive statistics for numeric columns")

doc.add_paragraph()
h2("3.2 Natural Language Query")
bullet("FR-07: ", "System shall accept free-text questions from the user via a chat input")
bullet("FR-08: ", "System shall maintain conversation history across multiple turns")
bullet("FR-09: ", "System shall use prior context when answering follow-up questions")
bullet("FR-10: ", "System shall support questions about aggregations, rankings, trends, comparisons, distributions, and correlations")

doc.add_paragraph()
h2("3.3 AI Code Generation")
bullet("FR-11: ", "System shall translate natural language questions into valid Pandas Python code")
bullet("FR-12: ", "System shall inject dataset schema and sample rows into the code generation prompt")
bullet("FR-13: ", "System shall attempt automatic error correction (up to 2 retries) if generated code fails")
bullet("FR-14: ", "Generated code shall always store its result in a variable named `result`")

doc.add_paragraph()
h2("3.4 Code Execution")
bullet("FR-15: ", "System shall execute generated code in an isolated sandbox")
bullet("FR-16: ", "System shall capture and display stdout output if present")
bullet("FR-17: ", "System shall return the result variable as the analysis output")
bullet("FR-18: ", "System shall display a user-friendly error message if execution fails after retries")

doc.add_paragraph()
h2("3.5 Visualization")
bullet("FR-19: ", "System shall automatically detect the most appropriate chart type")
bullet("FR-20: ", "System shall render line charts for time-series data")
bullet("FR-21: ", "System shall render bar charts for category comparisons and rankings")
bullet("FR-22: ", "System shall render histograms for distributions")
bullet("FR-23: ", "System shall render scatter plots for correlations")
bullet("FR-24: ", "System shall render pie charts for proportions/shares")
bullet("FR-25: ", "Charts shall be interactive (zoom, hover, download via Plotly)")

doc.add_paragraph()
h2("3.6 AI Explanation")
bullet("FR-26: ", "System shall generate a plain-English explanation for every analysis result")
bullet("FR-27: ", "Explanations shall directly answer the user's question in the first sentence")
bullet("FR-28: ", "Explanations shall highlight 2-3 key findings or patterns")
bullet("FR-29: ", "System shall suggest 3 contextual follow-up questions after each answer")
bullet("FR-30: ", "Follow-up questions shall be clickable and trigger new analyses")

doc.add_paragraph()
h2("3.7 Results Export")
bullet("FR-31: ", "System shall provide a download button for every tabular result")
bullet("FR-32: ", "Downloads shall be in CSV format")
bullet("FR-33: ", "Each download button shall be uniquely keyed to avoid UI conflicts")

doc.add_paragraph()
h2("3.8 Conversation Management")
bullet("FR-34: ", "System shall maintain a scrollable chat history within the session")
bullet("FR-35: ", "System shall provide a 'Clear Conversation' button to reset the session")
bullet("FR-36: ", "Conversation history shall include user messages, AI responses, charts, and data tables")

doc.add_page_break()

# ── 4. Non-Functional Requirements ───────────────────────────────────────────
h1("4. Non-Functional Requirements")

h2("4.1 Security")
bullet("NFR-01: ", "Generated code must pass a blocklist check before execution (blocks: subprocess, os, sys, socket, pickle, urllib, requests, __import__, eval, exec, open)")
bullet("NFR-02: ", "The executor must not allow filesystem reads or writes")
bullet("NFR-03: ", "The executor must not allow network requests")
bullet("NFR-04: ", "API keys must never be exposed in the UI or committed to version control")
bullet("NFR-05: ", "The .env file must be listed in .gitignore")

doc.add_paragraph()
h2("4.2 Performance")
bullet("NFR-06: ", "Code generation response time should be under 5 seconds for typical questions")
bullet("NFR-07: ", "Dataset upload and profiling should complete within 3 seconds for files up to 50MB")
bullet("NFR-08: ", "Chart rendering should be under 1 second for results with fewer than 10,000 rows")

doc.add_paragraph()
h2("4.3 Reliability")
bullet("NFR-09: ", "System shall auto-retry failed code generation up to 2 times before surfacing an error")
bullet("NFR-10: ", "System shall display user-friendly error messages — never raw Python tracebacks")
bullet("NFR-11: ", "System shall remain usable even if the AI explanation call fails")

doc.add_paragraph()
h2("4.4 Usability")
bullet("NFR-12: ", "UI shall be responsive and work on both desktop and tablet screen sizes")
bullet("NFR-13: ", "Dataset summary shall be visible in the sidebar at all times")
bullet("NFR-14: ", "The app shall provide example queries on the landing page")
bullet("NFR-15: ", "Every result shall have a visible Code tab so users can learn from the generated code")

doc.add_paragraph()
h2("4.5 Maintainability")
bullet("NFR-16: ", "Codebase shall follow a modular architecture with single-responsibility modules")
bullet("NFR-17: ", "AI model shall be configurable via a single MODEL constant in each module")
bullet("NFR-18: ", "API key shall be loaded from environment variables or Streamlit secrets — not hardcoded")

doc.add_page_break()

# ── 5. System Architecture ────────────────────────────────────────────────────
h1("5. System Architecture")

h2("5.1 High-Level Flow")
body("The system follows a linear pipeline for each user query:")
bullet("1. User uploads CSV → data_loader.py parses and profiles the dataset")
bullet("2. User submits a question → conversation.py records it in chat history")
bullet("3. query_engine.py sends question + dataset schema to Groq API")
bullet("4. Groq returns Pandas code → executor.py validates and runs it")
bullet("5. If execution fails → auto-correction loop (max 2 retries)")
bullet("6. visualization.py detects chart type and renders a Plotly figure")
bullet("7. ai_explainer.py sends result to Groq for plain-English explanation")
bullet("8. ai_explainer.py generates 3 suggested follow-up questions")
bullet("9. app.py renders the full response: explanation + chart + data + code + follow-ups")

doc.add_paragraph()
h2("5.2 Module Responsibilities")
add_table(
    ["Module", "Responsibility"],
    [
        ["app.py", "Streamlit UI, session state, main pipeline orchestration"],
        ["data_loader.py", "File parsing, datetime detection, dataset profiling"],
        ["query_engine.py", "Prompt construction, Groq API call, code extraction"],
        ["executor.py", "Safety validation, sandboxed exec(), result extraction"],
        ["visualization.py", "Chart type inference, Plotly figure creation"],
        ["ai_explainer.py", "Result explanation generation, follow-up question generation"],
        ["conversation.py", "Chat history management, context summarization"],
        ["utils.py", "CSV export, dtype labeling, number formatting"],
    ]
)

doc.add_paragraph()
h2("5.3 Technology Stack")
add_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["UI Framework", "Streamlit 1.35+", "Web interface, file upload, chat UI"],
        ["LLM Provider", "Groq API", "Fast inference for code generation & explanations"],
        ["LLM Model", "Llama 3.3 70B Versatile", "Code generation and natural language explanation"],
        ["Data Processing", "Pandas 2.0+ / NumPy 1.26+", "DataFrame operations and analysis"],
        ["Visualization", "Plotly 5.20+", "Interactive charts"],
        ["File Handling", "openpyxl 3.1+", "Excel file support"],
        ["Environment", "python-dotenv 1.0+", "Local API key management"],
        ["Deployment", "Streamlit Community Cloud", "Free public hosting"],
        ["Version Control", "GitHub", "Source code repository"],
    ]
)

doc.add_page_break()

# ── 6. Security Design ────────────────────────────────────────────────────────
h1("6. Security Design")

h2("6.1 Code Execution Sandbox")
body(
    "All AI-generated code is executed inside a controlled Python exec() environment. "
    "Before execution, the code is scanned against a regex blocklist. "
    "If any blocked pattern is found, execution is refused and an error is returned."
)
doc.add_paragraph()
body("Blocked patterns include:")
bullet("subprocess, shutil, socket, urllib, requests — network and process access")
bullet("pickle — arbitrary object deserialization")
bullet("__import__, importlib — dynamic imports")
bullet("import os, import sys — operating system access")

doc.add_paragraph()
h2("6.2 API Key Management")
body("Groq API keys are loaded from:")
bullet("Local: .env file (excluded from git via .gitignore)")
bullet("Cloud: Streamlit Secrets (encrypted at rest, injected as environment variables)")
body("Keys are never rendered in the UI or logged.")

doc.add_page_break()

# ── 7. Deployment ─────────────────────────────────────────────────────────────
h1("7. Deployment")

h2("7.1 Cloud Deployment (Streamlit Community Cloud)")
add_table(
    ["Property", "Value"],
    [
        ["Platform", "Streamlit Community Cloud"],
        ["Live URL", "https://chat-with-dataset-w3q3tiwx9b8duna8bdw9u6.streamlit.app"],
        ["GitHub Repo", "https://github.com/Himaaanshu7/chat-with-dataset"],
        ["Branch", "main"],
        ["Entry Point", "app.py"],
        ["Secret Required", "GROQ_API_KEY"],
        ["Cost", "Free"],
    ]
)

doc.add_paragraph()
h2("7.2 Local Setup")
body("Requirements: Python 3.10+, pip")
doc.add_paragraph()
body("Commands:")
p = doc.add_paragraph(style="No Spacing")
p.add_run(
    "git clone https://github.com/Himaaanshu7/chat-with-dataset.git\n"
    "cd chat-with-dataset\n"
    "pip install -r requirements.txt\n"
    "echo GROQ_API_KEY=your_key > .env\n"
    "streamlit run app.py"
).font.name = "Courier New"

doc.add_page_break()

# ── 8. Known Limitations ──────────────────────────────────────────────────────
h1("8. Known Limitations & Future Improvements")

h2("8.1 Current Limitations")
bullet("File size: Large files (>100MB) may be slow to profile")
bullet("LLM accuracy: Very complex multi-table joins may require rephrasing")
bullet("Session persistence: Conversation resets on page refresh")
bullet("Single file: Only one dataset can be active per session")

doc.add_paragraph()
h2("8.2 Planned Improvements")
bullet("SQL support: Allow users to query with SQL in addition to natural language")
bullet("Multiple file upload: Join and compare multiple datasets")
bullet("Export full report: Download entire conversation as PDF")
bullet("Authentication: User login and saved sessions")
bullet("Chart customization: Let users change chart type manually")
bullet("Database connectors: Connect directly to PostgreSQL, BigQuery, Snowflake")

doc.add_page_break()

# ── 9. Glossary ───────────────────────────────────────────────────────────────
h1("9. Glossary")
add_table(
    ["Term", "Definition"],
    [
        ["LLM", "Large Language Model — an AI model trained on text to generate code and language"],
        ["Groq", "AI inference platform providing fast LLM API access"],
        ["Llama 3.3 70B", "Open-source LLM by Meta with 70 billion parameters"],
        ["Pandas", "Python library for tabular data manipulation"],
        ["Plotly", "Python library for interactive charts"],
        ["Streamlit", "Python framework for building data web apps"],
        ["EDA", "Exploratory Data Analysis — the process of summarizing a dataset's key characteristics"],
        ["Sandbox", "Isolated execution environment that restricts dangerous operations"],
        ["Streamlit Secrets", "Encrypted key-value store for API keys in Streamlit Cloud deployments"],
    ]
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_page_break()
footer = doc.add_paragraph("Chat With Your Dataset — Project Documentation")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p2 = doc.add_paragraph(f"Himanshu Mishra · {datetime.date.today().strftime('%B %Y')}")
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.runs[0].font.size = Pt(9)
p2.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("Chat_With_Dataset_Documentation.docx")
print("Saved: Chat_With_Dataset_Documentation.docx")
